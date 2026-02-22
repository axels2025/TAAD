#!/usr/bin/env python3
"""Generate the TAAD Operations Manual as a professional .docx document.

Usage:
    python scripts/generate_operations_manual.py

Output:
    docs/TAAD_Operations_Manual_v1.0.docx
"""

import os
from datetime import date

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def set_cell_shading(cell, color_hex: str):
    """Set background shading on a table cell."""
    shading_elm = OxmlElement("w:shd")
    shading_elm.set(qn("w:fill"), color_hex)
    shading_elm.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_table(doc, headers: list[str], rows: list[list[str]], style="Table Grid"):
    """Add a formatted table to the document."""
    table = doc.add_table(rows=1, cols=len(headers), style=style)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr = table.rows[0]
    for i, header in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = header
        set_cell_shading(cell, "1a1a2e")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 212, 255)
                run.font.size = Pt(9)

    # Data rows
    for row_data in rows:
        row = table.add_row()
        for i, value in enumerate(row_data):
            cell = row.cells[i]
            cell.text = str(value)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph("")  # spacer
    return table


def add_code_block(doc, code: str):
    """Add a monospaced code block paragraph."""
    p = doc.add_paragraph()
    p.style = doc.styles["Code"]
    p.text = code
    return p


def add_bullet(doc, text: str, level: int = 0):
    """Add a bullet point."""
    p = doc.add_paragraph(text, style="List Bullet")
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.27 * (level + 1))
    return p


def add_numbered(doc, text: str):
    """Add a numbered list item."""
    return doc.add_paragraph(text, style="List Number")


def setup_styles(doc):
    """Configure document styles."""
    # Normal style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(6)

    # Code style
    if "Code" not in [s.name for s in doc.styles]:
        code_style = doc.styles.add_style("Code", WD_STYLE_TYPE.PARAGRAPH)
        code_style.font.name = "Consolas"
        code_style.font.size = Pt(8.5)
        code_style.font.color.rgb = RGBColor(30, 30, 30)
        code_style.paragraph_format.space_before = Pt(4)
        code_style.paragraph_format.space_after = Pt(4)
        pf = code_style.paragraph_format
        pf.left_indent = Cm(0.5)

    # Heading styles
    for level in range(1, 5):
        heading_style = doc.styles[f"Heading {level}"]
        heading_style.font.name = "Calibri"
        if level == 1:
            heading_style.font.size = Pt(18)
            heading_style.font.color.rgb = RGBColor(0, 100, 180)
        elif level == 2:
            heading_style.font.size = Pt(14)
            heading_style.font.color.rgb = RGBColor(0, 80, 150)
        elif level == 3:
            heading_style.font.size = Pt(12)
            heading_style.font.color.rgb = RGBColor(0, 60, 120)


def add_toc(doc):
    """Add a Table of Contents field."""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar = OxmlElement("w:fldChar")
    fldChar.set(qn("w:fldCharType"), "begin")
    run._r.append(fldChar)

    run2 = paragraph.add_run()
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
    run2._r.append(instrText)

    run3 = paragraph.add_run()
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    run3._r.append(fldChar2)

    run4 = paragraph.add_run("(Right-click and select 'Update Field' to generate Table of Contents)")
    run4.font.color.rgb = RGBColor(128, 128, 128)
    run4.font.italic = True

    run5 = paragraph.add_run()
    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "end")
    run5._r.append(fldChar3)


# ---------------------------------------------------------------------------
# Section writers
# ---------------------------------------------------------------------------

def write_title_page(doc):
    """Write the title page."""
    for _ in range(6):
        doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("TAAD")
    run.font.size = Pt(42)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 100, 180)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("Trade Archaeology & Alpha Discovery")
    run2.font.size = Pt(20)
    run2.font.color.rgb = RGBColor(0, 80, 150)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run("Operations Manual")
    run3.font.size = Pt(24)
    run3.font.bold = True

    doc.add_paragraph("")
    doc.add_paragraph("")

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = p4.add_run(f"Version 1.0  |  {date.today().strftime('%B %d, %Y')}")
    run4.font.size = Pt(12)
    run4.font.color.rgb = RGBColor(100, 100, 100)

    p5 = doc.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run5 = p5.add_run("Status: Production (All Phases Implemented)")
    run5.font.size = Pt(11)
    run5.font.color.rgb = RGBColor(0, 150, 0)

    doc.add_page_break()


def write_section_1(doc):
    """1. Executive Overview."""
    doc.add_heading("1. Executive Overview", level=1)

    doc.add_heading("What is TAAD?", level=2)
    doc.add_paragraph(
        "TAAD (Trade Archaeology & Alpha Discovery) is a self-learning agentic AI trading system "
        "that autonomously executes a proven naked put options strategy. It combines real-time market "
        "data from Interactive Brokers and Barchart with Claude AI reasoning to screen opportunities, "
        "score candidates, execute trades, monitor positions, and continuously learn from outcomes."
    )
    doc.add_paragraph(
        "The system implements a five-phase architecture: foundation infrastructure (Phase 0), "
        "baseline strategy encoding (Phase 1), autonomous execution (Phase 2), self-learning engine "
        "(Phase 3), AI-powered performance analysis (Phase 4), and a continuous agentic trading "
        "daemon (Phase 5). Each phase builds on the previous, creating a system that gets smarter "
        "with every trade."
    )
    doc.add_paragraph(
        "At its core, TAAD sells weekly out-of-the-money put options on liquid US equities and "
        "index products (SPX, XSP, SPY), collecting premium as income. It uses a six-dimension "
        "composite scoring algorithm to rank candidates, enforces strict multi-layered risk controls, "
        "and employs A/B experimentation to validate improvements before adopting them."
    )

    doc.add_heading("What Problem Does It Solve?", level=2)
    doc.add_paragraph(
        "Manual options trading is time-intensive, emotionally biased, and inconsistent. TAAD "
        "automates the entire weekly trading cycle: Sunday evening candidate screening, Monday "
        "morning execution, position monitoring throughout the week, and automated exits at profit "
        "targets or stop losses. The learning engine detects statistically significant patterns "
        "across 35+ dimensions, designs A/B experiments, and adapts parameters over time -- "
        "eliminating human cognitive biases while preserving the proven strategy's edge."
    )

    doc.add_heading("Who Is This Document For?", level=2)
    doc.add_paragraph(
        "This manual is for system operators running TAAD in paper or live trading, developers "
        "extending the system, and anyone evaluating the architecture. It assumes comfort with "
        "the command line and basic options trading concepts. A comprehensive glossary in Section 18 "
        "explains all technical terms."
    )

    doc.add_heading("Current System Status", level=2)
    add_table(doc,
        ["Phase", "Description", "Status"],
        [
            ["Phase 0", "Foundation (DB, IBKR, logging)", "Complete"],
            ["Phase 1", "Baseline Strategy Encoding", "Complete"],
            ["Phase 2", "Autonomous Execution", "Complete"],
            ["Phase 2.5", "Trade Data Collection (98 fields)", "Complete"],
            ["Phase 2.6", "Manual Trade Entry", "Complete"],
            ["Phase 3", "Self-Learning Engine (35+ patterns)", "Complete"],
            ["Phase 0-C", "Order Execution (160x faster)", "Complete"],
            ["Phase 4", "AI Insights Layer", "Complete"],
            ["Phase 5", "Continuous Agentic Daemon", "Complete"],
        ],
    )


def write_section_2(doc):
    """2. System Architecture."""
    doc.add_heading("2. System Architecture", level=1)

    doc.add_heading("High-Level Architecture", level=2)
    add_code_block(doc,
        "                        +-----------------------+\n"
        "                        |    TAAD Daemon (P5)   |\n"
        "                        |  Claude LLM Reasoning |\n"
        "                        |  Autonomy Governor    |\n"
        "                        +-----------+-----------+\n"
        "                                    |\n"
        "          +------------+------------+------------+\n"
        "          |            |            |            |\n"
        "    +-----+----+ +----+-----+ +----+----+ +-----+-----+\n"
        "    | Screening | | Scoring  | |Execution| | Learning  |\n"
        "    | Barchart  | | 6-Dim    | | IBKR    | | Patterns  |\n"
        "    | + IBKR    | | Composite| | Adaptive| | Experiments|\n"
        "    +-----+----+ +----+-----+ +----+----+ +-----+-----+\n"
        "          |            |            |            |\n"
        "          +------------+-----+------+------------+\n"
        "                             |\n"
        "                  +----------+----------+\n"
        "                  |   PostgreSQL / SQLite |\n"
        "                  |   18+ Tables          |\n"
        "                  |   Alembic Migrations   |\n"
        "                  +-----------------------+"
    )

    doc.add_heading("Data Flow", level=2)
    doc.add_paragraph(
        "The end-to-end data flow follows this pipeline:"
    )
    add_numbered(doc, "Screening: Barchart API scans entire US market -> ~50-100 candidates (1 API call, 2-3 seconds)")
    add_numbered(doc, "IBKR Validation: Real-time quotes, margin checks -> ~10-20 validated trades")
    add_numbered(doc, "Scoring: Six-dimension composite score (0-100) ranks each candidate")
    add_numbered(doc, "Staging: Top candidates staged with limit prices for next-day execution")
    add_numbered(doc, "Execution: Rapid-fire parallel orders via Adaptive Algo (<3 seconds for 5 orders)")
    add_numbered(doc, "Monitoring: 120-second interval position monitoring with Greeks and P&L")
    add_numbered(doc, "Exit Management: Profit targets (70%), stop losses (3x), time exits (3 DTE)")
    add_numbered(doc, "Snapshots: 98-field entry snapshot, daily position snapshots, 24-field exit snapshot")
    add_numbered(doc, "Learning: Pattern detection across 35+ dimensions, A/B experiments, parameter optimization")
    add_numbered(doc, "Reflection: Claude AI EOD analysis, weekly learning cycles, outcome feedback")

    doc.add_heading("Technology Stack", level=2)
    add_table(doc,
        ["Layer", "Technology", "Purpose"],
        [
            ["Language", "Python 3.11+", "Core runtime"],
            ["Database", "PostgreSQL (prod) / SQLite (dev)", "Data persistence"],
            ["ORM", "SQLAlchemy 2.0", "Database abstraction"],
            ["Migrations", "Alembic", "Schema management"],
            ["Broker API", "ib_insync 0.9.86", "Interactive Brokers integration"],
            ["Market Data", "Barchart onDemand API", "Screening and scanning"],
            ["AI/LLM", "Anthropic Claude (Opus/Sonnet)", "Reasoning engine, reflection"],
            ["ML", "scikit-learn, scipy", "Pattern detection, statistics"],
            ["CLI", "Typer + Rich", "Command-line interface"],
            ["Web Dashboard", "FastAPI + Uvicorn", "Monitoring dashboard"],
            ["Testing", "pytest (1790+ tests)", "Quality assurance"],
            ["Vector Search", "pgvector", "Semantic decision search"],
        ],
    )

    doc.add_heading("Directory Structure", level=2)
    add_code_block(doc,
        "trading_agent/\n"
        "+-- src/\n"
        "|   +-- agentic/           # Phase 5: Daemon, reasoning, autonomy\n"
        "|   +-- agents/            # AI agent base classes\n"
        "|   +-- cli/               # Typer CLI commands\n"
        "|   +-- config/            # Configuration management\n"
        "|   +-- data/              # Database models, migrations\n"
        "|   +-- execution/         # Order execution, exit management\n"
        "|   +-- learning/          # Pattern detection, experiments\n"
        "|   +-- nakedtrader/       # Daily SPX/XSP/SPY put selling\n"
        "|   +-- services/          # Market calendar, notifier, kill switch\n"
        "|   +-- strategies/        # Strategy definitions\n"
        "|   +-- tools/             # IBKR client, screener, data tools\n"
        "+-- config/                # YAML configs, systemd service\n"
        "+-- tests/                 # Unit, integration, E2E tests\n"
        "+-- scripts/               # Setup, seed data, utilities\n"
        "+-- data/                  # Databases, cache, exports\n"
        "+-- logs/                  # Application logs\n"
        "+-- docs/                  # Documentation"
    )


def write_section_3(doc):
    """3. Core Concepts & Trading Strategy."""
    doc.add_heading("3. Core Concepts & Trading Strategy", level=1)

    doc.add_heading("The Naked Put Strategy", level=2)
    doc.add_paragraph(
        "A naked put involves selling a put option without owning the underlying stock. The seller "
        "collects a premium upfront and profits if the stock stays above the strike price at expiration. "
        "The maximum profit is the premium received; the maximum loss is the strike price minus the "
        "premium (if the stock goes to zero). This strategy works best in stable or rising markets "
        "with liquid options and manageable implied volatility."
    )

    doc.add_heading("Key Strategy Parameters", level=2)
    add_table(doc,
        ["Parameter", "Value", "Description"],
        [
            ["Premium Target", "$0.30-$0.50", "Minimum premium per contract to collect"],
            ["Delta Range", "0.05-0.12 (NakedTrader)", "Target delta for strike selection"],
            ["OTM Range", "15-20%", "How far out-of-the-money"],
            ["DTE Range", "1-4 days (NakedTrader) / 7-14 (general)", "Days to expiration"],
            ["Contracts", "1-10", "Per trade"],
            ["Max Positions", "10", "Concurrent open positions"],
            ["Profit Target", "70% of max profit", "Exit at 70% premium decay"],
            ["Stop Loss", "3x premium (if enabled)", "Maximum loss per trade"],
            ["Time Exit", "3 DTE", "Close before expiration week risk"],
            ["Trend Filter", "Price > SMA20 > SMA50", "Only trade uptrends"],
            ["Max Margin Utilization", "80%", "Portfolio-level limit"],
            ["Max Sector Concentration", "30%", "Diversification requirement"],
        ],
    )

    doc.add_heading("The Options-First Approach", level=2)
    doc.add_paragraph(
        "Rather than screening stocks first and then checking options, TAAD uses a Barchart-first "
        "philosophy: scan the entire options market in a single API call, filter by premium, delta, "
        "DTE, and liquidity, then validate the top candidates with real-time IBKR data. This is "
        "10-20x faster than iterating through individual stock chains."
    )

    doc.add_heading("Supported Instruments", level=2)
    add_table(doc,
        ["Underlying", "Contract Type", "Settlement", "Assignment Style"],
        [
            ["SPX", "Index (SPXW)", "Cash", "European"],
            ["XSP", "Index (XSPW)", "Cash", "European"],
            ["SPY", "ETF", "Physical", "American"],
            ["~60 liquid equities", "Stock options", "Physical", "American"],
        ],
    )


def write_section_4(doc):
    """4. Installation & Setup."""
    doc.add_heading("4. Installation & Setup", level=1)

    doc.add_heading("Prerequisites", level=2)
    add_bullet(doc, "Python 3.11+ (required)")
    add_bullet(doc, "Interactive Brokers TWS or IB Gateway (paper trading account)")
    add_bullet(doc, "PostgreSQL 14+ (production) or SQLite (development)")
    add_bullet(doc, "Barchart onDemand subscription ($99/month for live, free trial for testing)")
    add_bullet(doc, "Anthropic API key (for Phase 4/5 AI features)")

    doc.add_heading("Environment Setup", level=2)
    add_code_block(doc,
        "# Clone the repository\n"
        "git clone <repo-url> trading_agent\n"
        "cd trading_agent\n"
        "\n"
        "# Create virtual environment\n"
        "python3.11 -m venv venv\n"
        "source venv/bin/activate\n"
        "\n"
        "# Install dependencies\n"
        "pip install -r requirements.txt\n"
        "\n"
        "# Configure environment\n"
        "cp .env.example .env\n"
        "# Edit .env with your API keys and settings"
    )

    doc.add_heading("Environment Variables (.env)", level=2)
    add_table(doc,
        ["Variable", "Default", "Description"],
        [
            ["IBKR_HOST", "127.0.0.1", "IBKR Gateway/TWS host"],
            ["IBKR_PORT", "7497", "7497 for paper, 7496 for live"],
            ["IBKR_CLIENT_ID", "1", "Default client ID"],
            ["IBKR_ACCOUNT", "DU123456", "Paper trading account ID"],
            ["DATABASE_URL", "sqlite:///data/databases/trades.db", "Database connection string"],
            ["ANTHROPIC_API_KEY", "", "Claude API key"],
            ["BARCHART_API_KEY", "", "Barchart onDemand API key"],
            ["PAPER_TRADING", "true", "Must be true for paper trading"],
            ["LOG_LEVEL", "INFO", "Logging verbosity"],
            ["USE_ADAPTIVE_ALGO", "true", "Enable Adaptive Algo orders"],
            ["ADAPTIVE_STRIKE_ENABLED", "true", "Delta-based strike selection"],
            ["STRIKE_TARGET_DELTA", "0.20", "Target absolute delta"],
            ["PREMIUM_FLOOR", "0.20", "Minimum acceptable premium"],
            ["MAX_DAILY_LOSS", "-0.02", "-2% circuit breaker"],
            ["LEARNING_ENABLED", "true", "Enable learning engine"],
        ],
    )

    doc.add_heading("Database Initialization", level=2)
    add_code_block(doc,
        "# Initialize database and run migrations\n"
        "python -m src.cli.main init\n"
        "alembic upgrade head\n"
        "\n"
        "# Verify setup\n"
        "python -m src.cli.main status\n"
        "python -m src.cli.main test-ibkr"
    )

    doc.add_heading("IBKR Configuration", level=2)
    doc.add_paragraph(
        "IBKR TWS or IB Gateway must be running with API connections enabled:"
    )
    add_bullet(doc, "Enable API: File > Global Configuration > API > Settings")
    add_bullet(doc, "Socket port: 7497 (paper trading)")
    add_bullet(doc, "Allow connections from localhost")
    add_bullet(doc, "Disable read-only API (for order placement)")

    doc.add_heading("IBKR Client ID Assignments", level=2)
    add_table(doc,
        ["Client ID", "Usage", "Notes"],
        [
            ["1", "Primary trading operations (scan, execute, trade, reconcile)", "Default"],
            ["2", "Position watch command", "Can run alongside ID 1"],
            ["3", "RESERVED - TWS Master API", "DO NOT USE"],
            ["4", "NakedTrader watch (nt-watch)", "Can run alongside others"],
            ["5", "NakedTrader sell (nt)", "Can run alongside watch commands"],
            ["10", "TAAD Daemon (Phase 5)", "Dedicated for autonomous daemon"],
        ],
    )


def write_section_5(doc):
    """5. Configuration."""
    doc.add_heading("5. Configuration", level=1)

    doc.add_heading("NakedTrader YAML Configuration", level=2)
    doc.add_paragraph("File: config/daily_spx_options.yaml")
    add_code_block(doc,
        "instrument:\n"
        "  default_symbol: XSP       # SPX, XSP, or SPY\n"
        "  contracts: 1               # Contracts per trade\n"
        "  max_contracts: 10\n"
        "\n"
        "strike:\n"
        "  delta_target: 0.065        # Target delta (6.5%)\n"
        "  delta_min: 0.05\n"
        "  delta_max: 0.12\n"
        "\n"
        "dte:\n"
        "  min: 1                     # Never 0 DTE\n"
        "  max: 4\n"
        "  prefer_shortest: true\n"
        "\n"
        "premium:\n"
        "  min_premium: 0.30\n"
        "\n"
        "exit:\n"
        "  profit_target_pct: 0.70    # Close at 70% profit\n"
        "  profit_target_floor: 0.10\n"
        "  stop_loss_enabled: false\n"
        "  stop_loss_multiplier: 3.0\n"
        "\n"
        "execution:\n"
        "  wait_for_open: true\n"
        "  open_delay_seconds: 60\n"
        "  order_type: LMT\n"
        "  fill_timeout_seconds: 300\n"
        "  latest_entry_time: '15:00'\n"
        "\n"
        "watch:\n"
        "  interval_seconds: 120\n"
        "  show_greeks: true"
    )

    doc.add_heading("Phase 5 Daemon Configuration", level=2)
    doc.add_paragraph("File: config/phase5.yaml")
    add_code_block(doc,
        "autonomy:\n"
        "  initial_level: 1           # Start at L1 (Recommend)\n"
        "  max_level: 2               # Max allowed autonomy\n"
        "  promotion_days: 14         # Clean days needed for promotion\n"
        "  promotion_min_trades: 20\n"
        "  promotion_min_win_rate: 0.70\n"
        "  demotion_loss_streak: 3\n"
        "\n"
        "claude:\n"
        "  reasoning_model: claude-opus-4-6\n"
        "  reflection_model: claude-sonnet-4-5-20250929\n"
        "  daily_cost_cap_usd: 10.00\n"
        "  max_reasoning_tokens: 4096\n"
        "  max_reflection_tokens: 2048\n"
        "\n"
        "daemon:\n"
        "  client_id: 10\n"
        "  heartbeat_interval_seconds: 60\n"
        "  event_poll_interval_seconds: 5\n"
        "  max_events_per_cycle: 10\n"
        "\n"
        "dashboard:\n"
        "  enabled: true\n"
        "  port: 8080\n"
        "  auth_token: ''             # Set for authentication"
    )

    doc.add_heading("Risk Governor Thresholds", level=2)
    add_table(doc,
        ["Limit", "Default", "Action When Exceeded"],
        [
            ["Daily Loss", "-2% of NLV", "HALT all trading"],
            ["Weekly Loss", "-5% of NLV", "HALT all trading"],
            ["Max Drawdown", "-10% from peak", "HALT all trading"],
            ["Max Positions", "10 concurrent", "Reject new trades"],
            ["Max Positions/Day", "10", "Reject new trades"],
            ["Max Margin Utilization", "80%", "Reject new trades"],
            ["Max Margin Per Trade", "50% of NLV", "Reject trade"],
            ["Sector Concentration", "30%", "Reject trade (min 4 positions)"],
            ["Position Stop Loss", "$500", "Exit position"],
            ["Min Excess Liquidity", "10% of NLV", "Margin danger alert"],
        ],
    )


def write_section_6(doc):
    """6. The Sunday-to-Monday Workflow."""
    doc.add_heading("6. The Sunday-to-Monday Workflow", level=1)

    doc.add_paragraph(
        "The primary operating procedure follows a weekly cycle: screen candidates on Sunday evening, "
        "execute trades Monday morning, and monitor positions throughout the week."
    )

    doc.add_heading("Sunday Session", level=2)

    doc.add_heading("Step 1: Barchart CSV Export", level=3)
    doc.add_paragraph(
        "Export a naked put screener CSV from Barchart.com with columns: Symbol, Price, Exp Date, "
        "DTE, Strike, Moneyness, Bid, Volume, Open Int, IV Rank, Delta, Return, Ann Rtn, Profit Prob."
    )

    doc.add_heading("Step 2: Run Sunday Session", level=3)
    add_code_block(doc,
        "python -m src.cli.main sunday-session --file naked-put-screener.csv"
    )
    doc.add_paragraph(
        "This command parses the CSV, connects to IBKR for live stock prices, validates candidates "
        "(OTM%, premium floor, DTE range, liquidity), calculates limit prices, determines contract "
        "quantities, and saves them as STAGED opportunities."
    )

    doc.add_heading("Step 3: Review Staged Trades", level=3)
    add_code_block(doc,
        "python -m src.cli.main show-staged     # View all staged\n"
        "python -m src.cli.main cancel-staged --id 3  # Remove unwanted"
    )

    doc.add_heading("Monday Morning Execution", level=2)

    doc.add_heading("9:15 AM ET - Pre-Market Validation", level=3)
    add_code_block(doc, "python -m src.cli.main validate-staged")
    doc.add_paragraph("Re-validates staged trades against fresh pre-market data. Flags stock price deviations.")

    doc.add_heading("9:28 AM ET - Quote Refresh", level=3)
    add_code_block(doc, "python -m src.cli.main refresh-quotes")
    doc.add_paragraph("Refreshes limit prices with live options quotes 2 minutes before market open.")

    doc.add_heading("9:30 AM ET - Execute", level=3)
    add_code_block(doc,
        "python -m src.cli.main execute-staged          # Standard execution\n"
        "python -m src.cli.main execute-two-tier --live  # Two-tier with monitoring"
    )
    doc.add_paragraph(
        "Orders are submitted via Adaptive Algo in parallel (<3 seconds for 5 orders). "
        "Fill monitoring runs for 10 minutes with progressive limit adjustments."
    )

    doc.add_heading("During the Week", level=2)
    add_bullet(doc, "Position monitoring: nakedtrader sell-watch (120-second intervals)")
    add_bullet(doc, "Daily position snapshots at 4:00 PM ET (automated via cron/systemd)")
    add_bullet(doc, "Exit management: 70% profit target, 3x stop loss, 3 DTE time exit")
    add_bullet(doc, "Reconciliation: python -m src.cli.main reconcile-positions")
    add_bullet(doc, "Order sync: python -m src.cli.main sync-orders")


def write_section_7(doc):
    """7. CLI Command Reference."""
    doc.add_heading("7. CLI Command Reference", level=1)

    # Infrastructure
    doc.add_heading("Infrastructure Commands", level=2)
    add_table(doc,
        ["Command", "Purpose", "Requires IBKR"],
        [
            ["init", "First-time setup (create DB, directories, logging)", "No"],
            ["status", "System health and trade statistics", "No"],
            ["test-ibkr", "Diagnose IBKR connection", "Yes"],
            ["version", "Show version and system info", "No"],
            ["db-reset", "Delete all data and recreate schema (destructive!)", "No"],
        ],
    )

    # Scanning
    doc.add_heading("Scanning Commands", level=2)
    add_table(doc,
        ["Command", "Key Flags", "Purpose"],
        [
            ["scan", "--max-results N, --no-validate, --save-file PATH", "Barchart API scan + IBKR validation"],
            ["scan --from-csv FILE", "--top N, --no-diversify, --show-rejected", "Import from Barchart CSV export"],
        ],
    )

    # Sunday-Monday Workflow
    doc.add_heading("Sunday-Monday Workflow Commands", level=2)
    add_table(doc,
        ["Command", "Key Flags", "Purpose"],
        [
            ["sunday-session", "--file PATH", "Stage trades from Barchart CSV"],
            ["show-staged", "", "View all staged opportunities"],
            ["cancel-staged", "--id N", "Remove a staged candidate"],
            ["validate-staged", "", "Re-validate with fresh market data"],
            ["refresh-quotes", "", "Update limit prices pre-execution"],
            ["execute-staged", "--dry-run / --live", "Execute staged orders"],
            ["execute-two-tier", "--mode hybrid|supervised|autonomous", "Two-tier execution with monitoring"],
        ],
    )

    # NakedTrader
    doc.add_heading("NakedTrader Commands", level=2)
    add_table(doc,
        ["Command", "Key Flags", "Purpose"],
        [
            ["nakedtrader sell SYMBOL", "--dry-run/--live, --yes, -c N, -d DELTA", "Sell a daily put option"],
            ["nakedtrader sell-watch", "--interval N, --once", "Monitor open NakedTrader positions"],
            ["nakedtrader sell-status", "--history N", "Show trade history and P&L summary"],
        ],
    )

    # Trade Management
    doc.add_heading("Trade Management Commands", level=2)
    add_table(doc,
        ["Command", "Key Flags", "Purpose"],
        [
            ["trade", "--auto, --max-trades N, --dry-run, --manual-only", "Full trade workflow"],
            ["execute SYMBOL STRIKE EXP", "--premium, --contracts, --dry-run", "Execute single trade"],
            ["add-trade", "--symbol, --strike, --expiration, --premium", "Add manual trade"],
            ["show-pending-trades", "--all-sources, --limit N", "View pending manual trades"],
            ["list-manual-trade-files", "--imported", "View manual trade JSON files"],
            ["web", "--host, --port", "Launch web interface for manual entry"],
            ["monitor", "", "Real-time position monitoring with P&L and Greeks"],
            ["list-trades", "--open-only", "Verify filled positions"],
        ],
    )

    # Order Sync
    doc.add_heading("Order Synchronization Commands", level=2)
    add_table(doc,
        ["Command", "Key Flags", "Purpose"],
        [
            ["sync-orders", "--date DATE, --include-filled", "Sync IBKR orders to database"],
            ["reconcile-positions", "--live (default: dry-run)", "Fix DB/IBKR discrepancies"],
        ],
    )

    # Learning
    doc.add_heading("Learning Commands", level=2)
    add_table(doc,
        ["Command", "Key Flags", "Purpose"],
        [
            ["learn --analyze", "", "Run weekly learning analysis"],
            ["learn --patterns", "", "View detected patterns"],
            ["learn --experiments", "", "View active A/B experiments"],
            ["learn --proposals", "", "View parameter proposals"],
            ["learn --summary", "--days N", "Summary of learning activity"],
            ["learning-stats", "", "Check data quality and coverage"],
            ["export-learning-data", "--output PATH", "Export learning data to CSV"],
            ["snapshot-positions", "", "Capture daily position snapshots"],
        ],
    )

    # TAAD Import
    doc.add_heading("TAAD Import Commands", level=2)
    add_table(doc,
        ["Command", "Key Flags", "Purpose"],
        [
            ["taad-import", "--account ID, --file PATH, --dry-run", "Import from IBKR Flex Query"],
            ["taad-status", "", "Check import session details"],
            ["taad-report", "--symbol, --sort pnl|symbol", "Matched trade P&L report"],
            ["taad-gaps", "", "Check data gaps and quality"],
            ["taad-promote", "--dry-run, --account ID", "Promote imports to Trade records"],
            ["taad-enrich", "--limit N, --with-scrape, --with-ibkr", "Enrich trades with market data"],
        ],
    )

    # Daemon
    doc.add_heading("Daemon Commands (Phase 5)", level=2)
    add_table(doc,
        ["Command", "Purpose"],
        [
            ["nakedtrader daemon start [--fg]", "Start the TAAD daemon"],
            ["nakedtrader daemon status", "Show daemon status, uptime, counters"],
            ["nakedtrader daemon context", "Show working memory state"],
            ["nakedtrader daemon pause", "Pause event processing"],
            ["nakedtrader daemon resume", "Resume from pause"],
            ["nakedtrader daemon set-autonomy LEVEL", "Set autonomy level (1-4)"],
            ["nakedtrader daemon audit [--limit N]", "Show decision audit log"],
            ["nakedtrader daemon costs", "Show Claude API cost summary"],
            ["nakedtrader daemon override ID [--action approve|reject]", "Override a pending decision"],
            ["nakedtrader daemon emergency-stop", "Emergency halt all trading"],
        ],
    )

    # Emergency
    doc.add_heading("Emergency Commands", level=2)
    add_table(doc,
        ["Command", "Key Flags", "Purpose"],
        [
            ["emergency-stop", "--liquidate", "Halt trading + optionally close all positions"],
            ["nakedtrader daemon emergency-stop", "", "Emergency stop via daemon CLI"],
        ],
    )


def write_section_8(doc):
    """8. Scoring Engine."""
    doc.add_heading("8. Scoring Engine", level=1)

    doc.add_paragraph(
        "Every candidate is scored on a 0-100 composite scale across six weighted dimensions. "
        "The score determines ranking for execution priority."
    )

    doc.add_heading("Six-Dimension Composite Score", level=2)
    add_table(doc,
        ["Dimension", "Weight", "Optimal Range", "What It Measures"],
        [
            ["Risk-Adjusted Return", "25%", "30-50% annualized", "Expected return relative to risk"],
            ["Probability of Profit", "20%", "80-90%", "Likelihood of expiring worthless (delta-based)"],
            ["IV Rank", "15%", "60-80%", "Current IV vs. 52-week range (premium richness)"],
            ["Liquidity", "15%", "OI >= 2000, Vol >= 300", "Ability to enter/exit without slippage"],
            ["Capital Efficiency", "15%", ">= 2.0% premium/strike", "Return per dollar of margin"],
            ["Safety Buffer", "10%", "12-18% OTM", "Distance from current price"],
        ],
    )

    doc.add_heading("Grade Assignment", level=2)
    add_table(doc,
        ["Score", "Grade", "Classification"],
        [
            ["85-100", "A+", "Excellent candidate"],
            ["75-84", "A", "Very good"],
            ["65-74", "B", "Good"],
            ["55-64", "C", "Acceptable"],
            ["45-54", "D", "Below average"],
            ["< 45", "F", "Poor -- skip"],
        ],
    )

    doc.add_heading("IBKR Validation Step", level=2)
    doc.add_paragraph(
        "After Barchart scoring, each candidate is validated against real-time IBKR data: "
        "live bid/ask confirms premium, whatIfOrder checks actual margin requirement, and "
        "contract qualification verifies tradability. Candidates failing validation are rejected "
        "with specific reasons."
    )


def write_section_9(doc):
    """9. Execution Engine."""
    doc.add_heading("9. Execution Engine", level=1)

    doc.add_heading("Order Executor", level=2)
    doc.add_paragraph(
        "Orders are placed through a multi-layer execution stack:"
    )
    add_bullet(doc, "Layer 1: IBKRClient -- connection management, automatic reconnection, audit logging")
    add_bullet(doc, "Layer 2: LiveStrikeSelector -- delta-based strike matching with fallback to OTM%")
    add_bullet(doc, "Layer 3: AdaptiveOrderExecutor -- IBKR Adaptive Algo (Urgent priority), LIMIT fallback")
    add_bullet(doc, "Layer 4: RapidFireExecutor -- parallel order submission (<3 seconds for 5 orders)")
    add_bullet(doc, "Layer 5: FillManager -- 10-minute monitoring, progressive $0.01/min adjustments, max 5 steps")
    add_bullet(doc, "Layer 6: OrderReconciliation -- sync database with TWS reality, discrepancy detection")

    doc.add_heading("Safety Verification", level=2)
    doc.add_paragraph(
        "On initialization, the order executor verifies PAPER_TRADING=true and IBKR_PORT=7497. "
        "It raises ValueError if either check fails, preventing accidental live trading."
    )

    doc.add_heading("Exit Manager", level=2)
    doc.add_paragraph("Four exit triggers evaluated in priority order:")
    add_table(doc,
        ["Exit Type", "Trigger", "Order Type", "Urgency"],
        [
            ["Profit Target", "P&L >= 70% of max profit", "LIMIT (1% buffer)", "Medium"],
            ["Stop Loss", "P&L <= -300% of premium", "MARKET (immediate)", "High"],
            ["Time Exit", "DTE <= 3 days", "LIMIT (2% buffer)", "Medium"],
            ["Stale Data", "No live market data", "SKIP (wait for data)", "N/A"],
        ],
    )

    doc.add_heading("Risk Governor", level=2)
    doc.add_paragraph("Nine pre-trade checks run sequentially before every order:")
    add_numbered(doc, "Trading halted? (kill switch check)")
    add_numbered(doc, "Duplicate position/order?")
    add_numbered(doc, "Earnings within DTE?")
    add_numbered(doc, "Daily loss limit exceeded? (-2% -> HALT)")
    add_numbered(doc, "Weekly loss limit exceeded? (-5% -> HALT)")
    add_numbered(doc, "Max drawdown exceeded? (-10% -> HALT)")
    add_numbered(doc, "Max positions reached? (10)")
    add_numbered(doc, "Max positions per day? (10)")
    add_numbered(doc, "Sector concentration too high? (>30%, min 4 positions)")


def write_section_10(doc):
    """10. Learning Engine."""
    doc.add_heading("10. Learning Engine", level=1)

    doc.add_heading("Architecture", level=2)
    add_code_block(doc,
        "LearningOrchestrator\n"
        "  +-- PatternDetector       # Finds patterns across 35+ dimensions\n"
        "  +-- PathAnalyzer          # Analyzes trade trajectories\n"
        "  +-- PatternCombiner       # Multi-dimensional combinations\n"
        "  +-- ExperimentEngine      # A/B testing framework\n"
        "  +-- ParameterOptimizer    # Proposes parameter changes"
    )

    doc.add_heading("Pattern Detection (35+ Patterns)", level=2)

    doc.add_heading("Technical Indicators (12 patterns)", level=3)
    doc.add_paragraph(
        "SMA alignment, EMA momentum, RSI oversold/overbought, MACD crossovers, "
        "Bollinger Band squeeze/touch, volume surge, proximity to 52-week high/low."
    )

    doc.add_heading("Market Context (7 patterns)", level=3)
    doc.add_paragraph(
        "VIX regime (low/medium/high), VIX term structure (contango/backwardation), "
        "SPY trend direction."
    )

    doc.add_heading("Trade Trajectories (9 patterns)", level=3)
    doc.add_paragraph(
        "Immediate winner, slow grind winner, early scare, steady loser, volatility spike, "
        "clean theta decay, delta blowout, IV crush, gamma risk realized."
    )

    doc.add_heading("Exit Quality (7 patterns)", level=3)
    doc.add_paragraph(
        "Early profit take, optimal exit, premature exit, necessary stop loss, premature stop loss, "
        "held to expiry, assignment risk."
    )

    doc.add_heading("Multi-Dimensional Combinations (10 patterns)", level=3)
    doc.add_paragraph(
        "Golden setup, high risk/high reward, safe income, earnings play, sector momentum, "
        "contrarian entry, momentum continuation, mean reversion, volatility expansion/contraction."
    )

    doc.add_heading("Statistical Validation", level=2)
    add_table(doc,
        ["Threshold", "Value", "Purpose"],
        [
            ["Minimum sample size", "30 trades per pattern", "Statistical validity"],
            ["P-value", "< 0.05", "Significance requirement"],
            ["Effect size", "> 0.5% ROI improvement", "Practical significance"],
            ["Confidence interval", "95%", "Statistical confidence"],
            ["Cross-validation", "5-fold", "Validation method"],
            ["Max parameter change", "1 per month", "Stability maintenance"],
            ["Max parameter shift", "20%", "Prevent over-tuning"],
        ],
    )

    doc.add_heading("A/B Experiment Engine", level=2)
    doc.add_paragraph(
        "When a significant pattern is detected, the experiment engine creates an A/B test: "
        "80% of trades use the control (current parameters), 20% use the test variant. "
        "After 30+ trades per group, a t-test determines whether to adopt (p < 0.05 and "
        "effect > 0.5% ROI), reject, or continue collecting data."
    )


def write_section_11(doc):
    """11. Data Collection & Snapshots."""
    doc.add_heading("11. Data Collection & Snapshots", level=1)

    doc.add_heading("Entry Snapshots (98 fields)", level=2)
    doc.add_paragraph(
        "Captured automatically at trade execution. Categories include:"
    )
    add_table(doc,
        ["Category", "Fields", "Examples"],
        [
            ["Option Data", "13", "symbol, strike, delta, gamma, theta, vega, IV, bid, ask"],
            ["Volatility", "5", "IV, IV rank, IV percentile, HV 20, IV/HV ratio"],
            ["Liquidity", "3", "option volume, open interest, volume/OI ratio"],
            ["Underlying", "6", "stock price, open, high, low, prev close, change %"],
            ["Calculated", "6", "OTM %, margin requirement, margin efficiency"],
            ["Trend", "6", "SMA 20/50, trend direction, trend strength"],
            ["Market", "4", "SPY price, VIX, SPY change %, VIX change %"],
            ["Events", "4", "earnings date, days to earnings, earnings in DTE"],
            ["Technical", "18", "RSI, MACD, ADX, ATR, Bollinger Bands, support/resistance"],
            ["Market Context", "10+", "QQQ, IWM, sector ETF, vol regime, market regime, FOMC"],
            ["Metadata", "4", "captured_at, source, data quality score"],
        ],
    )

    doc.add_heading("Daily Position Snapshots (15 fields)", level=2)
    doc.add_paragraph(
        "Captured at 4:00 PM ET via cron job or systemd timer. Tracks current bid/ask/mid, "
        "Greeks, underlying price, days in trade, unrealized P&L, max profit realized %, "
        "VIX, and underlying change."
    )
    add_code_block(doc,
        "# Schedule daily snapshots (cron)\n"
        "0 16 * * 1-5 TZ=America/New_York /path/to/scripts/daily_snapshot.sh\n"
        "\n"
        "# Or run manually\n"
        "python -m src.cli.main snapshot-positions"
    )

    doc.add_heading("Exit Snapshots (24 fields)", level=2)
    doc.add_paragraph(
        "Captured automatically when a trade closes. Records exit timestamp, exit price, "
        "exit type (profit_target/stop_loss/time_exit/manual), realized P&L, holding days, "
        "max drawdown, max profit during trade, underlying movement, VIX change, exit Greeks, "
        "and assignment status."
    )


def write_section_12(doc):
    """12. Database Schema."""
    doc.add_heading("12. Database Schema", level=1)

    doc.add_heading("Core Tables", level=2)
    add_table(doc,
        ["Table", "Purpose", "Key Fields"],
        [
            ["trades", "Complete trade lifecycle", "trade_id, symbol, strike, entry/exit dates, P&L, status"],
            ["positions", "Current open positions", "symbol, strike, contracts, unrealized P&L"],
            ["experiments", "A/B test tracking", "parameter, control/test values, p-value, decision"],
            ["learning_history", "What the system learned", "event_type, parameter changed, reasoning"],
            ["patterns", "Detected statistical patterns", "pattern_type, win_rate, confidence, p-value"],
            ["audit_log", "All AI decisions", "action, reasoning, confidence, outcome"],
            ["scan_results", "Scan sessions", "source, timestamp, candidate count"],
            ["scan_opportunities", "Individual candidates", "symbol, strike, score, status"],
            ["trade_entry_snapshots", "98-field entry capture", "All Greeks, market data, technicals"],
            ["position_snapshots", "Daily position state", "Current Greeks, P&L, underlying"],
            ["trade_exit_snapshots", "Exit capture", "Exit type, realized P&L, max drawdown"],
        ],
    )

    doc.add_heading("Phase 5 Daemon Tables", level=2)
    add_table(doc,
        ["Table", "Purpose", "Key Fields"],
        [
            ["daemon_events", "Durable event queue", "event_type, priority, status, payload"],
            ["decision_audit", "Full decision trail", "action, confidence, reasoning, autonomy gate, execution"],
            ["working_memory", "Crash-safe daemon state", "strategy_state, market_context, recent_decisions"],
            ["decision_embeddings", "Semantic search index", "text_content, vector (pgvector)"],
            ["daemon_health", "Heartbeat tracking", "pid, status, uptime, counters"],
            ["claude_api_costs", "API cost tracking", "model, tokens, cost_usd, daily_total"],
        ],
    )

    doc.add_heading("Migration Management", level=2)
    add_code_block(doc,
        "# Check current migration\n"
        "alembic current\n"
        "\n"
        "# Apply all pending migrations\n"
        "alembic upgrade head\n"
        "\n"
        "# Create new migration\n"
        "alembic revision --autogenerate -m \"description\"\n"
        "\n"
        "# Rollback one migration\n"
        "alembic downgrade -1"
    )


def write_section_13(doc):
    """13. External Integrations."""
    doc.add_heading("13. External Integrations", level=1)

    doc.add_heading("Interactive Brokers (IBKR)", level=2)
    doc.add_paragraph(
        "Connected via ib_insync library. Requires TWS or IB Gateway running locally. "
        "API capabilities used: order placement, real-time quotes, option chains, Greeks, "
        "whatIfOrder margin calculation, Adaptive Algo routing, Flex Query imports."
    )
    add_table(doc,
        ["Feature", "Method", "Notes"],
        [
            ["Order Placement", "ib.placeOrder()", "Adaptive Algo primary, LIMIT fallback"],
            ["Real-Time Quotes", "ib.reqMktData()", "Event-driven, <1 second latency"],
            ["Option Chains", "ib.reqSecDefOptParams()", "Full chain with all strikes/expirations"],
            ["Greeks", "ticker.modelGreeks", "Delta, gamma, theta, vega, IV"],
            ["Margin Check", "ib.whatIfOrder()", "Actual margin impact before execution"],
            ["Position Sync", "ib.positions()", "Reconciliation with database"],
            ["Fill Tracking", "ib.fills()", "Commission, fill price, fill time"],
        ],
    )

    doc.add_heading("Barchart", level=2)
    add_table(doc,
        ["Tier", "Cost", "Rate Limit", "Recommended For"],
        [
            ["Free Trial", "Free (30 days)", "400/day", "Development/testing"],
            ["onDemand Basic", "$99/month", "10,000/day", "Live trading (recommended)"],
            ["onDemand Professional", "$299/month", "50,000/day", "High-frequency scanning"],
        ],
    )

    doc.add_heading("Anthropic Claude API", level=2)
    doc.add_paragraph(
        "Used for AI reasoning (Phase 5 daemon) and performance analysis (Phase 4). "
        "Opus for complex reasoning decisions, Sonnet for EOD reflection. "
        "Daily cost cap enforced ($10/day default) with per-call cost tracking."
    )


def write_section_14(doc):
    """14. Risk Management."""
    doc.add_heading("14. Risk Management", level=1)

    doc.add_heading("Multi-Layered Risk Framework", level=2)

    doc.add_heading("Position-Level Controls", level=3)
    add_bullet(doc, "Delta limits: 0.05-0.12 for NakedTrader, 0.10-0.50 for general")
    add_bullet(doc, "OTM requirements: minimum 10-15% out of the money")
    add_bullet(doc, "DTE bounds: 1-4 days (NakedTrader), 7-14 days (general)")
    add_bullet(doc, "Premium floor: minimum $0.20-$0.30 per contract")
    add_bullet(doc, "Earnings avoidance: reject if earnings within DTE")

    doc.add_heading("Portfolio-Level Controls", level=3)
    add_bullet(doc, "Maximum 10 concurrent positions")
    add_bullet(doc, "80% maximum margin utilization")
    add_bullet(doc, "30% maximum sector concentration")
    add_bullet(doc, "50% maximum NLV per single trade margin")

    doc.add_heading("System-Level Controls", level=3)
    add_bullet(doc, "Daily loss circuit breaker: -2% of NLV halts all trading")
    add_bullet(doc, "Weekly loss limit: -5% halts all trading (resets Monday)")
    add_bullet(doc, "Maximum drawdown: -10% from peak equity halts all trading")
    add_bullet(doc, "Kill switch: persistent state, survives restarts")
    add_bullet(doc, "Emergency stop: CLI command immediately halts everything")

    doc.add_heading("Autonomy Controls (Phase 5)", level=3)
    add_table(doc,
        ["Level", "Name", "Capabilities", "Human Review"],
        [
            ["L1", "Recommend", "MONITOR_ONLY, REQUEST_HUMAN_REVIEW only", "All trades require approval"],
            ["L2", "Notify", "Routine actions if confidence >= 0.70", "Non-routine actions escalated"],
            ["L3", "Supervised", "Any action if confidence >= 0.60", "Mandatory triggers still escalate"],
            ["L4", "Autonomous", "Full autonomy", "9 mandatory triggers always escalate"],
        ],
    )

    doc.add_heading("9 Mandatory Escalation Triggers (All Levels)", level=3)
    add_numbered(doc, "Low confidence (< 0.60)")
    add_numbered(doc, "First trade of the day")
    add_numbered(doc, "New symbol (never traded before)")
    add_numbered(doc, "Loss exceeds threshold")
    add_numbered(doc, "Margin utilization > 60%")
    add_numbered(doc, "VIX spike (> 30 and > 20% increase)")
    add_numbered(doc, "Consecutive losses >= 3")
    add_numbered(doc, "Parameter change proposal")
    add_numbered(doc, "Stale market data")


def write_section_15(doc):
    """15. Monitoring & Operations."""
    doc.add_heading("15. Monitoring & Operations", level=1)

    doc.add_heading("Log Files", level=2)
    add_table(doc,
        ["File", "Purpose"],
        [
            ["logs/app.log", "General application log"],
            ["logs/trades.log", "Trade execution log"],
            ["logs/learning.log", "Learning engine log"],
            ["logs/errors.log", "Error-only log"],
            ["logs/daemon.log", "TAAD daemon output (systemd)"],
            ["logs/daemon-error.log", "TAAD daemon errors (systemd)"],
        ],
    )

    doc.add_heading("Daily Operational Checklist", level=2)
    add_numbered(doc, "Verify IBKR TWS/Gateway is running (test-ibkr)")
    add_numbered(doc, "Check system status (status command)")
    add_numbered(doc, "Review open positions (monitor or sell-watch)")
    add_numbered(doc, "Check for exit signals (profit targets, time exits)")
    add_numbered(doc, "Run position reconciliation (reconcile-positions)")
    add_numbered(doc, "Verify daily snapshots captured (learning-stats)")

    doc.add_heading("Weekly Review", level=2)
    add_numbered(doc, "Run learning analysis: learn --analyze")
    add_numbered(doc, "Review patterns: learn --patterns")
    add_numbered(doc, "Check experiments: learn --experiments")
    add_numbered(doc, "Review proposals: learn --proposals")
    add_numbered(doc, "Export learning data: export-learning-data")

    doc.add_heading("Troubleshooting", level=2)
    add_table(doc,
        ["Issue", "Diagnosis", "Fix"],
        [
            ["IBKR connection fails", "test-ibkr shows timeout", "Restart TWS/Gateway, check port 7497"],
            ["Stale quotes", "Monitor shows old prices", "Restart IBKR connection, check API permissions"],
            ["Migration errors", "alembic upgrade fails", "Check alembic current, review migration file"],
            ["Missing snapshots", "learning-stats shows gaps", "Verify cron job: crontab -l"],
            ["Orders not filling", "Status stays 'Submitted'", "Check market hours, spread width, premium"],
            ["Daemon not starting", "daemon start hangs", "Check PID file, kill stale process, review logs"],
            ["High Claude costs", "daemon costs shows spike", "Reduce daily_cost_cap_usd in phase5.yaml"],
        ],
    )

    doc.add_heading("Web Dashboard (Phase 5)", level=2)
    doc.add_paragraph(
        "The TAAD daemon serves a web dashboard at http://localhost:8080 with real-time "
        "status, open positions, recent decisions, pending approvals, and cost monitoring. "
        "Auto-refreshes every 10 seconds. Configure auth_token in phase5.yaml for security."
    )


def write_section_16(doc):
    """16. Testing."""
    doc.add_heading("16. Testing", level=1)

    doc.add_heading("Test Suite Overview", level=2)
    add_table(doc,
        ["Category", "Tests", "Coverage"],
        [
            ["Existing tests (Phases 0-4)", "1,790+", "87%+"],
            ["Phase 5 daemon tests", "380", "81-100%"],
            ["Total", "2,170+", "87%+"],
        ],
    )

    doc.add_heading("Running Tests", level=2)
    add_code_block(doc,
        "# All tests\n"
        "pytest\n"
        "\n"
        "# Unit tests only\n"
        "pytest tests/unit\n"
        "\n"
        "# Integration tests\n"
        "pytest tests/integration\n"
        "\n"
        "# With coverage report\n"
        "pytest --cov=src --cov-report=html\n"
        "\n"
        "# Specific Phase 5 tests\n"
        "pytest tests/unit/test_event_bus.py tests/unit/test_autonomy_governor.py\n"
        "\n"
        "# E2E daemon tests\n"
        "pytest tests/e2e/test_full_daemon_cycle.py"
    )

    doc.add_heading("Test Structure", level=2)
    add_bullet(doc, "tests/unit/ -- Unit tests for individual components (mocked dependencies)")
    add_bullet(doc, "tests/integration/ -- Integration tests for component interactions")
    add_bullet(doc, "tests/e2e/ -- End-to-end tests simulating full daemon cycles")
    add_bullet(doc, "tests/conftest.py -- Shared fixtures (in-memory SQLite, temp databases)")


def write_section_17(doc):
    """17. Development & Extension."""
    doc.add_heading("17. Development & Extension", level=1)

    doc.add_heading("Code Standards", level=2)
    add_bullet(doc, "PEP 8 compliance with 88-character line length (Black default)")
    add_bullet(doc, "Type hints required on all function signatures")
    add_bullet(doc, "Google-style docstrings for all public functions/classes")
    add_bullet(doc, "Imports: stdlib, third-party, local (three groups)")

    add_code_block(doc,
        "# Code quality commands\n"
        "black src/ tests/          # Auto-format\n"
        "ruff check src/ tests/     # Lint\n"
        "mypy src/                  # Type check"
    )

    doc.add_heading("Adding a New Pattern Detector", level=2)
    add_numbered(doc, "Add pattern buckets to src/learning/pattern_detector.py")
    add_numbered(doc, "Define the field source (TradeEntrySnapshot column)")
    add_numbered(doc, "Add bucket boundaries and names")
    add_numbered(doc, "The statistical framework (t-test, confidence) applies automatically")
    add_numbered(doc, "Add unit tests in tests/unit/")

    doc.add_heading("Adding a New CLI Command", level=2)
    add_numbered(doc, "Create function in appropriate commands module under src/cli/commands/")
    add_numbered(doc, "Decorate with @app.command() or add to Typer subgroup")
    add_numbered(doc, "Register in src/cli/main.py if new subgroup")
    add_numbered(doc, "Add tests")

    doc.add_heading("Database Migration Workflow", level=2)
    add_code_block(doc,
        "# 1. Modify models in src/data/models.py\n"
        "# 2. Generate migration\n"
        "alembic revision --autogenerate -m \"Add new_column to trades\"\n"
        "# 3. Review generated migration in src/data/migrations/versions/\n"
        "# 4. Apply\n"
        "alembic upgrade head\n"
        "# 5. Test\n"
        "pytest tests/"
    )


def write_section_18(doc):
    """18. Glossary."""
    doc.add_heading("18. Glossary", level=1)

    glossary = [
        ("A/B Testing (Experiment Engine)", "A method of comparing two variants (control vs. test) to determine which performs better. TAAD allocates 80% of trades to the current strategy (control) and 20% to a proposed improvement (test), then uses statistical tests to decide whether to adopt the change."),
        ("Adaptive Algo", "An IBKR smart order routing algorithm that dynamically navigates the bid-ask spread to optimize fill prices. TAAD uses Adaptive Algo with 'Urgent' priority as the primary order type, falling back to standard LIMIT orders when unavailable."),
        ("Alembic", "A database migration tool for SQLAlchemy. It tracks schema changes over time and can upgrade or downgrade the database to any version. Run 'alembic upgrade head' to apply all pending migrations."),
        ("API (Application Programming Interface)", "A set of rules that allows software programs to communicate. TAAD uses APIs from Interactive Brokers (trading), Barchart (screening), and Anthropic (AI reasoning)."),
        ("Assignment", "When an option seller is obligated to fulfill the contract. For put sellers, this means buying the stock at the strike price. European-style options (SPX, XSP) can only be assigned at expiration; American-style (SPY, stocks) can be assigned at any time."),
        ("ATM (At-the-Money)", "An option with a strike price equal to or very close to the current stock price. TAAD targets OTM puts, not ATM."),
        ("Autonomy Level", "Phase 5 daemon's permission tier (L1-L4). L1 (Recommend) requires human approval for all trades. L4 (Autonomous) can execute independently, except for 9 mandatory escalation triggers."),
        ("Barchart", "A market data provider. TAAD uses Barchart's options screener API to scan the entire US market in a single API call, then validates candidates with IBKR."),
        ("Bid/Ask Spread", "The difference between the highest price a buyer will pay (bid) and the lowest price a seller will accept (ask). Narrow spreads indicate good liquidity."),
        ("Black (Formatter)", "A Python code formatter that enforces consistent style. Run 'black src/ tests/' to auto-format all code."),
        ("Bollinger Bands", "A technical indicator consisting of a moving average and two standard deviation bands. Used by the learning engine to detect squeeze and touch patterns."),
        ("Bracket Order", "An order group consisting of a parent order (sell put) and two child orders (profit target buy and stop loss buy). When the parent fills, children activate as OCA (one-cancels-all)."),
        ("CBOE PutWrite Index", "A benchmark index that tracks the performance of a hypothetical portfolio selling cash-secured S&P 500 put options. Used as a reference for strategy performance."),
        ("Circuit Breaker", "An automatic safety mechanism that halts all trading when a loss threshold is reached. TAAD has three: daily (-2%), weekly (-5%), and max drawdown (-10%)."),
        ("CLI (Command Line Interface)", "A text-based interface for interacting with software. TAAD's CLI is built with Typer and provides 40+ commands for trading, monitoring, and management."),
        ("Contango", "A VIX term structure where near-term VIX futures are cheaper than longer-term futures. This is the normal state and generally favorable for put selling."),
        ("Credit Spread", "A multi-leg options strategy involving buying and selling options at different strikes. Not currently used by TAAD, which focuses on naked puts."),
        ("Cron Job", "A scheduled task on Unix/Linux/macOS systems. TAAD uses cron for daily position snapshots at 4:00 PM ET."),
        ("CSV (Comma-Separated Values)", "A simple file format for tabular data. Barchart exports screening results as CSV files that TAAD imports via the sunday-session command."),
        ("Delta", "The rate of change of an option's price relative to a $1 change in the underlying stock. For puts, delta is negative. A delta of -0.10 means the option loses $0.10 for every $1 the stock rises. TAAD targets deltas of 0.05-0.12 (NakedTrader)."),
        ("Drawdown", "The peak-to-trough decline in portfolio value, expressed as a percentage. TAAD's maximum drawdown limit is -10%."),
        ("DTE (Days to Expiration)", "The number of calendar days until an option contract expires. TAAD's NakedTrader targets 1-4 DTE."),
        ("Effect Size", "A statistical measure of the practical significance of a result. TAAD requires an effect size > 0.5% ROI improvement before adopting a pattern or parameter change."),
        ("Entry Snapshot", "A comprehensive 98-field data capture taken at the moment of trade execution. Includes option data, Greeks, underlying data, technical indicators, market context, and metadata."),
        ("European-Style Option", "An option that can only be exercised at expiration. SPX and XSP options are European-style, eliminating early assignment risk."),
        ("Exit Snapshot", "A 24-field data capture taken when a trade closes. Records exit type, realized P&L, holding days, max drawdown during trade, and market conditions at exit."),
        ("Flex Query", "An IBKR reporting tool that generates XML files with detailed execution and trade data. TAAD's taad-import command processes Flex Query exports."),
        ("Gamma", "The rate of change of delta per $1 change in the underlying. High gamma means delta can change rapidly, increasing risk as expiration approaches."),
        ("Historical Volatility (HV)", "A measure of how much a stock's price has moved in the past, typically over 20 trading days. Compared to implied volatility to assess option pricing."),
        ("IB Gateway", "A lightweight, headless version of TWS for API-only access. Recommended for production use as it doesn't require a display."),
        ("ib_insync", "A Python library that provides a synchronous, Pythonic interface to the Interactive Brokers API. TAAD uses it for all broker interactions."),
        ("Implied Volatility (IV)", "The market's expectation of future price movement, embedded in the option's price. Higher IV means higher premiums for sellers."),
        ("Iron Condor", "A four-leg options strategy. Not used by TAAD, which focuses on single-leg naked puts."),
        ("ITM (In-the-Money)", "A put option with a strike price above the current stock price. TAAD avoids ITM options."),
        ("IV Percentile", "The percentage of days in the past year when IV was lower than the current level. High IV percentile suggests options are relatively expensive."),
        ("IV Rank", "Current IV relative to its 52-week high and low, expressed as a percentage. An IV Rank of 60% means current IV is 60% of the way between its annual low and high. TAAD scores higher in the 60-80% range."),
        ("JSON (JavaScript Object Notation)", "A lightweight data format used for configuration and data exchange. TAAD stores structured data (payloads, snapshots) as JSON in PostgreSQL columns."),
        ("JSONB", "PostgreSQL's binary JSON format, which supports indexing and efficient querying. Used for daemon event payloads and working memory storage."),
        ("Kill Switch", "A persistent safety mechanism in TAAD that halts all trading. Survives restarts (file-based state). Activated by emergency-stop command."),
        ("Learning Orchestrator", "The top-level coordinator for TAAD's learning engine. It runs pattern detection, experiment evaluation, and parameter optimization in a structured weekly cycle."),
        ("Limit Order", "An order to buy or sell at a specified price or better. TAAD primarily uses limit orders for put selling, with the Adaptive Algo navigating toward the best fill."),
        ("Margin", "The collateral required by the broker to maintain a short option position. For naked puts, margin is typically 15-20% of the notional value."),
        ("Margin Efficiency", "The ratio of premium collected to margin required. Higher values mean more income per dollar of capital committed. TAAD targets 1:10 to 1:20 ratios."),
        ("Market Order", "An order to buy or sell at the current market price. TAAD uses market orders only for urgent stop-loss exits."),
        ("Migration", "A versioned database schema change managed by Alembic. Each migration file describes how to upgrade (and downgrade) the database."),
        ("mypy", "A static type checker for Python. Verifies that type hints are consistent and correct."),
        ("Naked Put", "Selling a put option without owning the underlying stock or having an offsetting position. The seller collects premium and profits if the stock stays above the strike price."),
        ("Net Liquidation Value (NLV)", "The total value of all positions in an account if they were liquidated immediately. Used as the denominator for percentage-based risk limits."),
        ("ORM (Object-Relational Mapping)", "A technique that maps database tables to Python objects. TAAD uses SQLAlchemy as its ORM."),
        ("OTM (Out-of-the-Money)", "A put option with a strike price below the current stock price. TAAD targets 15-20% OTM puts for safety."),
        ("Paper Trading", "Simulated trading with fake money to test strategies without financial risk. TAAD should run in paper trading mode for 3-6 months before considering live trading."),
        ("Path Analyzer", "A component of the learning engine that analyzes how trades evolve over time (trajectory patterns like 'immediate winner' or 'early scare')."),
        ("Pattern Combiner", "Combines multiple individual patterns (e.g., low delta + high IV rank + uptrend) into multi-dimensional patterns for more nuanced strategy insights."),
        ("Pattern Detector", "The primary learning component that analyzes trade history across 35+ dimensions to find statistically significant patterns."),
        ("PEP 8", "Python's official style guide for code formatting. TAAD follows PEP 8 with an 88-character line length."),
        ("pgvector", "A PostgreSQL extension for vector similarity search. TAAD uses it to store and query embeddings of past decisions for semantic similarity."),
        ("pip", "Python's package installer. Run 'pip install -r requirements.txt' to install all dependencies."),
        ("Position Snapshot", "A 15-field daily capture of each open position's current state (Greeks, P&L, underlying price). Used by the path analyzer in the learning engine."),
        ("PostgreSQL", "An open-source relational database. TAAD uses PostgreSQL for production deployments (SQLite for development/testing)."),
        ("Premium", "The price received when selling an option. This is the maximum profit for a naked put seller. TAAD targets $0.30-$0.50 per contract."),
        ("pytest", "Python's most popular testing framework. TAAD has 2,170+ tests across unit, integration, and E2E categories."),
        ("Rich", "A Python library for beautiful terminal output with colors, tables, and progress bars. Used by TAAD's CLI for display formatting."),
        ("Risk Governor", "TAAD's multi-layered risk management system that checks 9+ conditions before every trade, including daily loss limits, margin utilization, and sector concentration."),
        ("ROI (Return on Investment)", "Profit as a percentage of capital invested. For options, typically calculated as premium profit / margin required."),
        ("RSI (Relative Strength Index)", "A momentum indicator ranging 0-100. Below 30 is 'oversold,' above 70 is 'overbought.' Used by the pattern detector."),
        ("Ruff", "A fast Python linter. Run 'ruff check src/ tests/' to check for code quality issues."),
        ("Scoring Engine", "TAAD's six-dimension composite scoring algorithm (0-100) that evaluates candidates across risk-adjusted return, probability of profit, IV rank, liquidity, capital efficiency, and safety buffer."),
        ("Sharpe Ratio", "A measure of risk-adjusted return: (return - risk-free rate) / standard deviation. Higher is better. TAAD targets a Sharpe Ratio of 1.5+."),
        ("SQLAlchemy", "Python's most popular ORM and database toolkit. TAAD uses it for all database operations with both PostgreSQL and SQLite backends."),
        ("Strike Price", "The price at which the option can be exercised. For puts, the seller may be obligated to buy the stock at this price."),
        ("Sunday-to-Monday Workflow", "TAAD's primary operating procedure: screen candidates Sunday evening, validate Monday morning, execute at market open, monitor through the week."),
        ("systemd", "A Linux service manager. TAAD includes a systemd service file (config/taad.service) for running the daemon as a managed service."),
        ("TAAD", "Trade Archaeology & Alpha Discovery. The name of the overall trading system. The 'Archaeology' refers to mining historical trade data for insights; 'Alpha Discovery' refers to finding market edges."),
        ("Theta", "The rate at which an option loses value over time (time decay). Positive for option sellers. TAAD profits from theta decay as options approach expiration."),
        ("Trade Archaeology", "The process of importing, analyzing, and learning from historical IBKR trades via Flex Query exports. Implemented in the taad-import command pipeline."),
        ("TWS (Trader Workstation)", "Interactive Brokers' desktop trading application. Provides both manual trading interface and API access."),
        ("Type Hints", "Python annotations that specify expected types for function arguments and return values. Required by TAAD's coding standards."),
        ("Typer", "A Python library for building CLI applications with type hints. TAAD's entire command-line interface is built on Typer."),
        ("Variance Risk Premium", "The difference between implied volatility and realized volatility. When IV > HV, option sellers collect a premium for bearing volatility risk."),
        ("Vega", "The sensitivity of an option's price to a 1% change in implied volatility. Important for understanding how IV changes affect position value."),
        ("Virtual Environment", "An isolated Python installation that keeps project dependencies separate. Created with 'python -m venv venv' and activated with 'source venv/bin/activate'."),
        ("VIX", "The CBOE Volatility Index, measuring expected 30-day S&P 500 volatility. Often called the 'fear gauge.' VIX > 25 triggers caution in TAAD; VIX > 30 triggers mandatory escalation."),
        ("Win Rate", "The percentage of trades that are profitable. TAAD's baseline target is 75%+."),
    ]

    for term, definition in glossary:
        p = doc.add_paragraph()
        run_term = p.add_run(term + ": ")
        run_term.bold = True
        run_term.font.size = Pt(10)
        p.add_run(definition).font.size = Pt(10)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    """Generate the full Operations Manual."""
    doc = Document()
    setup_styles(doc)

    # Title page
    write_title_page(doc)

    # Table of Contents
    doc.add_heading("Table of Contents", level=1)
    add_toc(doc)
    doc.add_page_break()

    # Sections
    write_section_1(doc)
    doc.add_page_break()
    write_section_2(doc)
    doc.add_page_break()
    write_section_3(doc)
    doc.add_page_break()
    write_section_4(doc)
    doc.add_page_break()
    write_section_5(doc)
    doc.add_page_break()
    write_section_6(doc)
    doc.add_page_break()
    write_section_7(doc)
    doc.add_page_break()
    write_section_8(doc)
    doc.add_page_break()
    write_section_9(doc)
    doc.add_page_break()
    write_section_10(doc)
    doc.add_page_break()
    write_section_11(doc)
    doc.add_page_break()
    write_section_12(doc)
    doc.add_page_break()
    write_section_13(doc)
    doc.add_page_break()
    write_section_14(doc)
    doc.add_page_break()
    write_section_15(doc)
    doc.add_page_break()
    write_section_16(doc)
    doc.add_page_break()
    write_section_17(doc)
    doc.add_page_break()
    write_section_18(doc)

    # Save
    output_dir = os.path.join(os.path.dirname(__file__), "..", "docs")
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, "TAAD_Operations_Manual_v1.0.docx")
    doc.save(output_path)
    print(f"Operations manual saved to: {output_path}")
    print(f"File size: {os.path.getsize(output_path) / 1024:.1f} KB")
    print("Note: Open in Word and right-click the Table of Contents to 'Update Field' for page numbers.")


if __name__ == "__main__":
    main()
